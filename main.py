import tkinter as tk
from tkinter import messagebox, StringVar, Text, WORD, BOTH, END, Scrollbar, ttk, filedialog
import logging
from email_sender import EmailSender
from template_processor import TemplateProcessor
from rules import RuleManager
from gui.app_window import AppWindow
import os
from datetime import datetime
import webbrowser
from docx import Document
import sys

class MergeMailApp:
    def __init__(self):
        # Tạo root window trước
        self.root = tk.Tk()
        
        # Sau đó mới tạo các StringVar
        self.csv_path = StringVar(self.root)
        self.template_path = StringVar(self.root)
        
        # Khởi tạo các components
        self.email_sender = EmailSender()
        self.template_processor = TemplateProcessor()
        self.rule_manager = RuleManager()
        
        # Tạo window chính và đóng root window
        self.window = AppWindow(self)
        self.root.withdraw()

        # Thêm xử lý khi đóng cửa sổ
        self.window.window.protocol("WM_DELETE_WINDOW", self.on_closing)

    def on_closing(self):
        """Xử lý khi đóng cửa sổ"""
        if messagebox.askokcancel("Thoát", "Bạn có chắc muốn thoát?"):
            self.window.window.destroy()
            self.root.quit()
            sys.exit(0)  # Đảm bảo thoát hoàn toàn

    def run(self):
        self.window.run()

    def browse_csv(self):
        filename = filedialog.askopenfilename(
            title="Chọn file CSV",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if filename:
            self.csv_path.set(filename)
            self.window.progress_frame.update_status(f"Đã chọn file CSV: {os.path.basename(filename)}")
            return True
        return False

    def browse_template(self):
        filename = filedialog.askopenfilename(
            title="Chọn file Template",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.template_path.set(filename)
            self.window.progress_frame.update_status(f"Đã chọn file template: {os.path.basename(filename)}")
            return True
        return False

    def start_merge_mail(self):
        try:
            if not self.csv_path.get() or not self.template_path.get():
                messagebox.showerror("Error", "Vui lòng chọn cả file CSV và template")
                return

            # Lấy thông tin email, password và subject
            email_frame = self.window.get_email_frame()
            email = email_frame.get_email()
            password = email_frame.get_password()
            subject_template = email_frame.get_subject()  # Lấy template tiêu đề

            if not email or not password:
                messagebox.showerror("Error", "Vui lòng nhập email và App Password")
                return

            if not subject_template:
                messagebox.showerror("Error", "Vui lòng nhập tiêu đề email")
                return

            try:
                data = self.template_processor.read_csv(self.csv_path.get())
                self.window.progress_frame.update_status("Đã đọc file CSV")
            except Exception as e:
                messagebox.showerror("Error", f"Lỗi đọc file CSV: {str(e)}")
                return

            try:
                template = self.template_processor.read_template(self.template_path.get())
                self.window.progress_frame.update_status("Đã đọc file template")
            except Exception as e:
                messagebox.showerror("Error", f"Lỗi đọc file template: {str(e)}")
                return

            try:
                self.email_sender.configure(email, password)
                self.window.progress_frame.update_status("Đã cấu hình email")
            except Exception as e:
                messagebox.showerror("Error", f"Lỗi cấu hình email: {str(e)}")
                return

            # Thêm progress bar
            total_rows = len(data)
            success_count = 0
            error_count = 0
            skipped_count = 0

            logging.info(f"Bắt đầu xử lý {total_rows} email")

            for index, row in data.iterrows():
                try:
                    if self.rule_manager.check_rules(row):
                        self.window.progress_frame.update_status(f"Đang xử lý email {index + 1}/{total_rows}")
                        
                        # Tạo nội dung email
                        merged_content = self.template_processor.merge_template(template, row)
                        logging.info(f"Đã merge template cho {row['email']}")
                        
                        # Tạo tiêu đề email với placeholder
                        merged_subject = self.template_processor.merge_template(
                            {'content': subject_template, 'images': []}, 
                            row
                        )
                        subject = merged_subject['content']
                        
                        # Gửi email
                        self.email_sender.send_email(
                            to_email=row['email'],
                            subject=subject,
                            content=merged_content
                        )
                        success_count += 1
                        logging.info(f"Gửi thành công đến {row['email']}")
                    else:
                        skipped_count += 1
                        logging.warning(f"Bỏ qua email {row['email']} do không đạt rules")
                    
                    # Cập nhật progress
                    progress = ((index + 1) / total_rows) * 100
                    self.window.progress_frame.update_progress(progress)
                    
                except Exception as e:
                    error_count += 1
                    logging.error(f"Lỗi gửi email đến {row.get('email', 'unknown')}: {str(e)}")
                    self.window.progress_frame.update_status(f"Lỗi: {str(e)}")
                    continue

            # Hiển thị kết quả chi tiết
            result_message = (
                f"Hoàn thành!\n"
                f"Tổng số: {total_rows}\n"
                f"Đã gửi: {success_count}\n"
                f"Bỏ qua: {skipped_count}\n"
                f"Lỗi: {error_count}"
            )
            messagebox.showinfo("Kết quả", result_message)
            
            self.window.progress_frame.update_status("Sẵn sàng gửi email")
            self.window.progress_frame.update_progress(0)
            
        except Exception as e:
            messagebox.showerror("Error", str(e))
            logging.error(f"Lỗi trong quá trình gửi email: {str(e)}")
            self.window.progress_frame.update_status("Đã xảy ra lỗi")

    def show_guide(self):
        # Tạo cửa sổ popup
        guide_window = tk.Toplevel(self.window.window)
        guide_window.title("Hướng dẫn Sử Dụng")
        guide_window.geometry("800x600")
        
        # Tạo frame chứa nội dung
        main_frame = ttk.Frame(guide_window, padding="20")
        main_frame.pack(fill=BOTH, expand=True)

        # Tạo widget Text có scroll
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=BOTH, expand=True)
        
        text_widget = Text(
            text_frame,
            wrap=WORD,
            padx=10,
            pady=10,
            font=("Helvetica", 10),
            background="white",
            foreground="black"
        )
        
        # Thêm scrollbar
        scrollbar = Scrollbar(text_frame)
        scrollbar.pack(side="right", fill="y")
        text_widget.pack(side="left", fill=BOTH, expand=True)
        
        # Kết nối scrollbar với text widget
        text_widget.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=text_widget.yview)

        # Nội dung hướng dẫn
        guide_text = """
HƯỚNG DẪN SỬ DỤNG ỨNG DỤNG MERGE MAIL

I. CHUẨN BỊ
-----------

1. Cấu hình Gmail:
   • Bật xác thực 2 lớp cho Gmail
   • Tạo App Password:
     - Vào Google Account > Security > 2-Step Verification
     - Tạo App Password và lưu lại mã 16 ký tự

2. Chuẩn bị File:
   a) File CSV (danh sách người nhận):
      • Định dạng UTF-8
      • Phải có các cột: name, email
      • Ví dụ:
        name,email,company,position
        Nguyễn Văn A,nguyenvana@gmail.com,Công ty ABC,Giám đốc
        
   b) File Template Word (.docx):
      • Sử dụng placeholder: {{ tên_cột }}
      • Ví dụ:
        Kính gửi {{ name }},
        Chúng tôi gửi thư này đến {{ position }} của {{ company }}.

II. CÁCH SỬ DỤNG
---------------

1. Chọn File:
   • Click "Browse CSV" để chọn file danh sách
   • Click "Browse Template" để chọn file mẫu email

2. Nhập Thông Tin Gmail:
   • Gmail: nhập địa chỉ email của bạn
   • App Password: nhập mã 16 ký tự đã tạo

3. Gửi Email:
   • Click "Start Merge and Send"
   • Theo dõi tiến trình trên thanh progress

III. TÍNH NĂNG BỔ SUNG
--------------------

1. Tạo File Mẫu:
   • Click "Tạo file mẫu"
   • File được tạo trong thư mục "samples"

2. Xử Lý Lỗi:
   • Tự động thử lại khi gặp lỗi (tối đa 3 lần)
   • Log lỗi được lưu trong error_log.txt

IV. LƯU Ý QUAN TRỌNG
------------------

1. Kiểm tra kỹ:
   • Nội dung template trước khi gửi
   • Định dạng file CSV (UTF-8)
   • Tên các cột trong CSV

2. An Toàn:
   • Không chia sẻ App Password
   • Test với số lượng nhỏ trước
   • Kiểm tra kết nối mạng

3. Hỗ Trợ:
   • Lỗi được ghi vào error_log.txt
   • Liên hệ admin nếu cần hỗ trợ thêm
"""

        # Chèn nội dung và cấu hình
        text_widget.insert(END, guide_text)
        text_widget.config(state="disabled")  # Không cho phép chỉnh sửa

        # Nút đóng
        close_button = ttk.Button(
            main_frame,
            text="Đóng",
            command=guide_window.destroy,
            style='Accent.TButton'
        )
        close_button.pack(pady=10)

        # Đặt cửa sổ ở giữa màn hình
        guide_window.update_idletasks()
        width = guide_window.winfo_width()
        height = guide_window.winfo_height()
        x = (guide_window.winfo_screenwidth() // 2) - (width // 2)
        y = (guide_window.winfo_screenheight() // 2) - (height // 2)
        guide_window.geometry(f'{width}x{height}+{x}+{y}')
        
        # Đặt focus cho cửa sổ mới
        guide_window.focus_set()
        guide_window.grab_set()  # Modal window

    def create_sample_files(self):
        # Tạo file mẫu
        try:
            if not os.path.exists("samples"):
                os.makedirs("samples")
            
            # Tạo file CSV mẫu
            csv_content = """name,email,company,position
Nguyễn Văn A,nguyenvana@example.com,Công ty ABC,Giám đốc
Trần Thị B,tranthib@example.com,Công ty XYZ,Trưởng phòng
Lê Văn C,levanc@example.com,Công ty DEF,Nhân viên"""
            
            csv_path = os.path.join("samples", f"sample_data_{datetime.now().strftime('%Y%m%d')}.csv")
            with open(csv_path, "w", encoding="utf-8") as f:
                f.write(csv_content)
            
            # Tạo file Word mẫu
            doc = Document()
            doc.add_paragraph("Kính gửi {{ name }},")
            doc.add_paragraph("\nChúng tôi gửi thư này đến {{ position }} của {{ company }}.")
            doc.add_paragraph("\nĐây là email tự động được gửi từ hệ thống merge mail.")
            doc.add_paragraph("\nTrân trọng,\nBan Quản trị")
            
            doc_path = os.path.join("samples", f"sample_template_{datetime.now().strftime('%Y%m%d')}.docx")
            doc.save(doc_path)
            
            messagebox.showinfo(
                "Thành công", 
                f"Đã tạo file mẫu trong thư mục 'samples':\n"
                f"1. {os.path.basename(csv_path)}\n"
                f"2. {os.path.basename(doc_path)}"
            )
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tạo file mẫu: {str(e)}")

    def open_help(self):
        webbrowser.open('https://myaccount.google.com/apppasswords')

    def browse_attachment(self):
        """Chọn file đính kèm"""
        filenames = filedialog.askopenfilenames(
            title="Chọn file đính kèm",
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"),
                ("All files", "*.*")
            ]
        )
        if filenames:
            try:
                for filename in filenames:
                    self.email_sender.add_attachment(filename)
                
                # Cập nhật label
                files_text = "\n".join([os.path.basename(f) for f in filenames])
                self.window.get_email_frame().update_attachment_label(files_text)
                self.window.progress_frame.update_status(f"Đã thêm {len(filenames)} file đính kèm")
            except Exception as e:
                messagebox.showerror("Lỗi", str(e))

    def clear_attachments(self):
        """Xóa tất cả file đính kèm"""
        self.email_sender.clear_attachments()
        self.window.get_email_frame().update_attachment_label("Chưa có file nào được chọn")
        self.window.progress_frame.update_status("Đã xóa tất cả file đính kèm")

if __name__ == "__main__":
    app = MergeMailApp()
    app.run() 